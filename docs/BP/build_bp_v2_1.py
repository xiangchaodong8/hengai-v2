# -*- coding: utf-8 -*-
"""Merge terminology + brand appendix into Co2Lion_BP_2026_v2.docx -> v2.1.docx (styled)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

BASE = Path(__file__).resolve().parent
SRC = BASE / "Co2Lion_BP_2026_v2.docx"
OUT = BASE / "Co2Lion_BP_2026_v2.2.docx"

FONT = "Arial"
C_CHAPTER = RGBColor(0x1B, 0x3A, 0x5C)
C_SUB = RGBColor(0x0B, 0x6E, 0x4F)
C_BODY = RGBColor(0x33, 0x33, 0x33)

OLD_TM = (
    '"Co2Lion"、"方舟魔盒"、"CL-GTS"、"CL-IVC"、"HengAI"、"西值东送，全球大同"'
    "及相关标识为本公司商标或正在申请注册的商标，未经授权不得使用。"
)
NEW_TM = (
    '"Co2Lion"、"方舟魔盒"、"CL-GTS"、"CL-IVC"、"CL-MAT"、"HengAI"、'
    '"GreenMark（绿印）"、"西值东送，全球大同"及相关标识为本公司商标或正在申请注册的商标，'
    "未经授权不得使用。"
)

GLOSSARY = {
    "（一）政策法规类": [
        ("双碳", "—", "中国2030年前碳达峰、2060年前碳中和的国家级气候承诺；Co2Lion政策契合论述的根本依据。"),
        ("西电东送", "—", "西部清洁能源经特高压输送至东部；GEC+与「西值东送，全球大同」品牌图腾的物理基础。"),
        ("LCA", "Life Cycle Assessment", "生命周期评估；CBAM嵌入排放核算的方法论基础。"),
        ("GHG Protocol", "Greenhouse Gas Protocol", "WRI/WBCSD温室气体核算体系；ISO 14067、CBAM与Co2Lion数据质量的重要参照。"),
        ("CBAM", "Carbon Border Adjustment Mechanism", "欧盟碳边界调节机制；2026起实质付费阶段。"),
        ("EU ETS", "EU Emissions Trading System", "欧盟碳排放交易体系；CBAM碳价挂钩基准（如€75/tCO₂e）。"),
        ("十五五", "—", "2026—2030国家规划期；碳足迹管理体系、绿色制造等政策窗口。"),
        ("人工智能+", "AI+", "国发〔2025〕11号文及工信部八部门专项行动；HengAI第二条政策护城河。"),
        ("数据二十条", "—", "数据要素市场化配置政策框架；供应链碳数据确权与跨境合规路径。"),
        ("专精特新", "—", "建议申报「小巨人」：工业软件与碳数据合规方向；政府背书与上市通道价值。"),
    ],
    "（二）技术协议类（Co2Lion自研）": [
        ("CL-GTS", "Co2Lion Green Trade Stack", "绿色贸易技术栈；Protocol v16.0 Full Stack。"),
        ("CL-IVC", "Co2Lion Industrial Verification Chain", "工业隐私/主权盾；ZKP「数据可用不可见、物理不出域」。"),
        ("CL-MAT / 方舟魔盒", "MAT Edge Gateway", "边缘感知网关；Lv.4物理采集；DMC样机已完成。"),
        ("GEC", "Green Energy Certificate", "绿电证书；GEC+之前置概念。"),
        ("GEC+", "GEC Plus", "方舟关税碳盾；西电东送绿色信用数字化路由。"),
        ("Lv.4 数据置信度", "Data Confidence Level 4", "Lv.1人工→Lv.4 MAT传感器；欧盟CBAM免审核最高档。"),
        ("ZKP", "Zero-Knowledge Proof", "零知识证明；供应商合规验证而不泄露BOM。"),
        ("ETS 时空平移", "ETS Time-Space Shift", "储能充放电时移减排认证；ETS/VCM可流通碳资产。"),
        ("因子精算", "Factor Authority Network", "工业原厂1—9号工序碳强度精算与CL-IVC确权；产业链公共因子基础设施。"),
        ("NHJC", "能耗在线监测技术规范", "重点用能单位监测国标（NHJC-04/06/08-2018）；CL-MAT架构对齐，非等同已通过验收。"),
        ("COP", "Carbon Ownership Protocol", "碳所有权协议；与UNFCCC COP形成「责任/资产」平行呼应。"),
        ("COS", "Carbon Ownership Summit", "年度碳所有权峰会；三大榜单+社群仪式。"),
        ("ESC 能算耦合账本", "—", "ETS时空平移的底层记账逻辑（BP商业化扩展层）。"),
        ("产业链主权先行登记", "—", "HengAI荣誉机制：发起升格确权即获轻量认可，永久有效、不因流程中断撤销。"),
    ],
    "（三）金融与商业类": [
        ("ARR", "Annual Recurring Revenue", "年度经常性收入；3年目标人民币4.5亿元。"),
        ("TAM / SAM / SOM", "—", "可触达/可服务/可获得市场；BP叁市场规模分层。"),
        ("Green-Fi", "—", "绿色金融数据路由；银行风控与碳资产映射。"),
        ("VCM", "Voluntary Carbon Market", "自愿碳市场；碳信用开发服务层。"),
        ("VPP", "Virtual Power Plant", "虚拟电厂。"),
        ("ESG", "Environmental, Social, Governance", "环境、社会与治理；国际买家与金融机构合规要求背景。"),
        ("链主", "Supply Chain Anchor", "掌握终端市场准入的头部出口企业；供应链协同付费与网络效应触发源。"),
        ("GM / 绿印", "GreenMark", "HengAI行为积分；非财政承诺或投资收益。"),
        ("Pre-money / 出让比例", "—", "本轮3,000万；Pre-money 8—12亿；出让2.4%—3.6%。"),
    ],
    "（四）行业机构与平台类": [
        ("ISO 14067", "—", "产品碳足迹国际标准；Co2Lion参与技术工作。"),
        ("CISA / 中钢协", "China Iron and Steel Association", "钢铁因子精算报表样式兼容对象。"),
        ("CBAM Registry", "—", "欧盟CBAM官方申报平台。"),
        ("授权申报人", "EU Authorized CBAM Declarant", "欧盟CBAM正式申报主体资质；Pre-A报关行收购目标。"),
        ("BOM", "Bill of Materials", "物料清单；CL-IVC核心保护对象。"),
        ("Core Climate", "—", "香港碳信用国际化通道。"),
        ("DMC 样机", "Design Manufacturing Certification Prototype", "设计制造认证样机；方舟魔盒量产前验证阶段。"),
        ("APAK", "—", "国际亮相节点（里程碑2027 Q2）。"),
        ("HengAI / 全域中心", "HengAI Hub", "V2商业版智能工作台；14业务模块+产业主权看板等。"),
        ("能效领跑者", "—", "工信部等联合公示名单；与Co2Lion八大行业重合的可核实获客清单。"),
    ],
}

TRADEMARKS = [
    ("公司", "Co2Lion、方舟数字科技"),
    ("产品/协议", "方舟魔盒、CL-GTS、CL-IVC、CL-MAT、HengAI"),
    ("积分品牌", "GreenMark（绿印）"),
    ("口号/图腾", "西值东送，全球大同"),
]

BRAND_RULES = [
    ("主标识", "Co2Lion拉丁字标+中文「方舟数字科技」；封面以字标为主，与正文四层架构图分级呈现。"),
    ("产品标识", "HengAI、CL-IVC、CL-GTS、方舟魔盒在架构图中使用，不与公司主标识争抢视觉中心。"),
    ("色彩", "主绿（合规/增长）+金色（主权/价值）；与Lv.4确权、绿印积分语义一致。"),
    ("合规禁用", "不得将「西值东送」与财政补贴比例、政府背书承诺绑在同一视觉块。"),
]

PRODUCT_ALIGN = [
    ("工业原厂·因子精算", "原厂因子精算模块；精算真理源，非看板内嵌。"),
    ("产业主权看板", "全域中心 #origin-audit；驾驶舱壳层，不算力真理源。"),
    ("供应链协同", "全域中心 #supply；穿透填报与共振协同。"),
    ("决策层呈送包", "全域中心 #decision；企业内部呈批，非CL代写财政申请。"),
    ("碳税敞口", "默认€为主口径；UI可切换¥参考折算（非银行牌价）。"),
]


def fmt_def(term: str, en: str, definition: str) -> str:
    if en and en != "—":
        if en == term or term in en:
            return definition
        return f"{en}。{definition}"
    return definition


def style_run(run, *, bold=False, size=11, color=C_BODY, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color


def style_paragraph(p, *, align=None, space_before=None, space_after=None):
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)


def insert_paragraph_before(anchor: Paragraph, text: str = "", kind: str = "body"):
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    para = Paragraph(new_p, anchor._parent)
    if not text:
        return para

    run = para.add_run(text)
    if kind == "chapter":
        style_run(run, bold=True, size=18, color=C_CHAPTER)
        style_paragraph(para, space_before=24, space_after=10)
    elif kind == "subsection":
        style_run(run, bold=True, size=14, color=C_SUB)
        style_paragraph(para, space_before=16, space_after=7)
    else:
        style_run(run, size=11, color=C_BODY)
        style_paragraph(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=4, space_after=5)
    return para


def set_cell_text(cell, text: str, *, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    style_paragraph(p, space_before=2, space_after=2)
    run = p.add_run(text)
    style_run(run, bold=bold, size=10, color=C_BODY)


def insert_table_before(anchor: Paragraph, headers, rows):
    doc = anchor.part.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Normal Table"
    except KeyError:
        pass

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri + 1].cells[ci], val)

    anchor._p.addprevious(table._tbl)
    return table


def build_appendix_blocks():
    blocks = [
        ("chapter", "拾贰  术语表"),
        (
            "body",
            "以下术语表供投资人快速查阅，分类与正文肆—拾壹各章呼应。"
            "英文缩写首次出现处，以「英文全称。中文释义」格式呈现。",
        ),
    ]

    for section, items in GLOSSARY.items():
        blocks.append(("subsection", section))
        rows = [(term, fmt_def(term, en, defn)) for term, en, defn in items]
        blocks.append(("table", ["术语", "释义"], rows))

    blocks.extend(
        [
            ("subsection", "品牌标识规范"),
            (
                "body",
                "注册商标与视觉组合与本文封面、封底及第五章知识产权声明保持一致；"
                "印刷物料不得擅自改字或混用未授权标识。",
            ),
            (
                "table",
                ["类别", "标识"],
                TRADEMARKS,
            ),
            ("subsection", "LOGO使用原则"),
            ("table", ["项目", "说明"], BRAND_RULES),
            ("subsection", "HengAI V2产品术语对齐"),
            ("table", ["BP用语", "说明"], PRODUCT_ALIGN),
        ]
    )
    return blocks


def update_trademark(doc: Document):
    for p in doc.paragraphs:
        if OLD_TM in p.text:
            p.text = p.text.replace(OLD_TM, NEW_TM)
            return True
    return False


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)

    doc = Document(str(SRC))
    updated = update_trademark(doc)

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "数字文明，铸就大国承诺":
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("Cannot find back-cover anchor paragraph")

    for block in reversed(build_appendix_blocks()):
        kind = block[0]
        if kind in ("chapter", "subsection", "body"):
            insert_paragraph_before(anchor, block[1], kind=kind)
        elif kind == "table":
            insert_table_before(anchor, block[1], block[2])

    doc.save(str(OUT))
    term_count = sum(len(v) for v in GLOSSARY.values())
    print(f"OK: {OUT.name}")
    print(f"  trademark_updated={updated}")
    print(f"  glossary_terms={term_count}")
    print(f"  style=Arial/chapter+subsection+2col-table")


if __name__ == "__main__":
    main()
